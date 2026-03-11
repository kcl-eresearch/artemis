"""Output format writers for Artemis research results.

Provides functions for writing research essays as JSON, Markdown, or DOCX.
Used by both the research CLI (cli.py) and the standalone converter (convert.py).
"""

import json
import re
import sys
from typing import Any


def format_sources_md(results: list) -> str:
    """Build a numbered Markdown sources list from search results.

    Args:
        results: List of objects with .title and .url attributes,
                 or dicts with "title" and "url" keys.
    """
    if not results:
        return ""
    lines = ["\n\n---\n\n## Sources\n"]
    for i, r in enumerate(results, 1):
        if isinstance(r, dict):
            title = r.get("title") or r.get("url", "")
            url = r.get("url", "")
        else:
            title = r.title or r.url
            url = r.url
        lines.append(f"{i}. [{title}]({url})")
    return "\n".join(lines)


def write_json(
    path: str,
    query: str,
    essay: str,
    results: list,
    usage: dict[str, Any] | None = None,
    *,
    stdout: bool = False,
) -> None:
    """Write results as JSON."""
    output = {
        "query": query,
        "essay": essay,
        "results": [
            (
                {"title": r.get("title", ""), "url": r.get("url", ""), "snippet": r.get("snippet", "")}
                if isinstance(r, dict)
                else {"title": r.title, "url": r.url, "snippet": r.snippet}
            )
            for r in results
        ],
        "usage": usage,
    }
    if stdout:
        json.dump(output, sys.stdout, indent=2)
        print()
    else:
        with open(path, "w") as f:
            json.dump(output, f, indent=2)


def write_markdown(
    path: str,
    essay: str,
    results: list | None = None,
    *,
    stdout: bool = False,
) -> None:
    """Write essay as Markdown with optional sources appendix."""
    content = essay
    if results:
        content += format_sources_md(results)
    if stdout:
        print(content)
    else:
        with open(path, "w") as f:
            f.write(content)


def _add_hyperlink(paragraph, url: str, text: str):
    """Add a clickable, blue-underlined hyperlink to a python-docx paragraph."""
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement
    from docx.shared import Pt, RGBColor

    part = paragraph.part
    r_id = part.relate_to(
        url,
        "http://schemas.openxmlformats.org/officeDocument/2006/relationships/hyperlink",
        is_external=True,
    )

    hyperlink = OxmlElement("w:hyperlink")
    hyperlink.set(qn("r:id"), r_id)

    run = OxmlElement("w:r")
    rPr = OxmlElement("w:rPr")

    # Blue colour
    color = OxmlElement("w:color")
    color.set(qn("w:val"), "0563C1")
    rPr.append(color)

    # Underline
    u = OxmlElement("w:u")
    u.set(qn("w:val"), "single")
    rPr.append(u)

    # Font size matching body text
    sz = OxmlElement("w:sz")
    sz.set(qn("w:val"), str(Pt(11).pt * 2))  # half-points
    rPr.append(sz)

    run.append(rPr)

    t = OxmlElement("w:t")
    t.text = text
    run.append(t)
    hyperlink.append(run)
    paragraph._p.append(hyperlink)
    return hyperlink


def md_to_docx(
    path: str,
    essay: str,
    title: str | None = None,
    results: list | None = None,
) -> None:
    """Convert a Markdown essay into a formatted DOCX document.

    Args:
        path: Output file path
        essay: Markdown essay text
        title: Optional document title (added as heading level 0)
        results: Optional list of source results for a sources appendix.
                 Each item should have .title/.url attrs or "title"/"url" keys.
    """
    from docx import Document
    from docx.enum.text import WD_ALIGN_PARAGRAPH

    doc = Document()

    if title:
        title_para = doc.add_heading(title, level=0)
        title_para.alignment = WD_ALIGN_PARAGRAPH.CENTER

    for line in essay.split("\n"):
        stripped = line.strip()
        if not stripped:
            doc.add_paragraph("")
            continue

        if stripped.startswith("#### "):
            doc.add_heading(stripped[5:], level=4)
        elif stripped.startswith("### "):
            doc.add_heading(stripped[4:], level=3)
        elif stripped.startswith("## "):
            doc.add_heading(stripped[3:], level=2)
        elif stripped.startswith("# "):
            doc.add_heading(stripped[2:], level=1)
        elif stripped.startswith("- ") or stripped.startswith("* "):
            doc.add_paragraph(stripped[2:], style="List Bullet")
        elif re.match(r"^\d+\.\s", stripped):
            text = re.sub(r"^\d+\.\s", "", stripped)
            doc.add_paragraph(text, style="List Number")
        else:
            doc.add_paragraph(stripped)

    if results:
        doc.add_page_break()
        doc.add_heading("Sources", level=1)
        for i, r in enumerate(results, 1):
            if isinstance(r, dict):
                r_title = r.get("title") or r.get("url", "")
                r_url = r.get("url", "")
            else:
                r_title = r.title or r.url
                r_url = r.url
            para = doc.add_paragraph(f"{i}. ")
            _add_hyperlink(para, r_url, r_title)

    doc.save(path)
