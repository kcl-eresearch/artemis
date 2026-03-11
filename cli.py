#!/usr/bin/env python3
"""One-shot CLI for deep research without the HTTP server.

Usage examples::

    # Basic research (outputs JSON by default)
    python cli.py "quantum computing advances"

    # DOCX report with auto-generated filename
    python cli.py "climate change mitigation" --format docx

    # Markdown to specific file, shallow preset
    python cli.py "rust vs go" --format md --output comparison.md --preset shallow

    # Write to stdout
    python cli.py "LLM architectures" --format md --output -
"""

import asyncio
import argparse
import json
import re
import sys
from datetime import date

from dotenv import load_dotenv

load_dotenv()

_FORMATS = ("json", "md", "docx")
_PRESETS = ("deep", "shallow")


def _slugify(text: str, max_len: int = 48) -> str:
    """Convert a query string into a filename-safe slug."""
    slug = text.lower().strip()
    slug = re.sub(r"[^\w\s-]", "", slug)
    slug = re.sub(r"[\s_]+", "-", slug)
    slug = re.sub(r"-+", "-", slug).strip("-")
    return slug[:max_len]


def _default_output_path(query: str, fmt: str) -> str:
    """Generate an output filename from query and date."""
    ext = {"json": "json", "md": "md", "docx": "docx"}[fmt]
    slug = _slugify(query)
    today = date.today().isoformat()
    return f"{slug}-{today}.{ext}"


def _format_sources(results: list) -> str:
    """Build a numbered sources list from search results."""
    if not results:
        return ""
    lines = ["\n\n---\n\n## Sources\n"]
    for i, r in enumerate(results, 1):
        title = r.title or r.url
        lines.append(f"{i}. [{title}]({r.url})")
    return "\n".join(lines)


def _write_json(path: str, query: str, result, *, stdout: bool = False) -> None:
    """Write results as JSON."""
    output = {
        "query": query,
        "essay": result.essay,
        "results": [
            {"title": r.title, "url": r.url, "snippet": r.snippet}
            for r in result.results
        ],
        "usage": result.usage.model_dump() if result.usage else None,
    }
    if stdout:
        json.dump(output, sys.stdout, indent=2)
        print()
    else:
        with open(path, "w") as f:
            json.dump(output, f, indent=2)


def _write_markdown(path: str, query: str, result, *, stdout: bool = False) -> None:
    """Write essay as Markdown with a sources appendix."""
    content = result.essay + _format_sources(result.results)
    if stdout:
        print(content)
    else:
        with open(path, "w") as f:
            f.write(content)


def _add_hyperlink(paragraph, url: str, text: str):
    """Add a clickable hyperlink to a python-docx paragraph."""
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement

    part = paragraph.part
    r_id = part.relate_to(url, "http://schemas.openxmlformats.org/officeDocument/2006/relationships/hyperlink", is_external=True)

    hyperlink = OxmlElement("w:hyperlink")
    hyperlink.set(qn("r:id"), r_id)

    r = OxmlElement("w:r")
    rPr = OxmlElement("w:rPr")
    rStyle = OxmlElement("w:rStyle")
    rStyle.set(qn("w:val"), "Hyperlink")
    rPr.append(rStyle)
    r.append(rPr)

    t = OxmlElement("w:t")
    t.text = text
    r.append(t)
    hyperlink.append(r)
    paragraph._p.append(hyperlink)
    return hyperlink


def _write_docx(path: str, query: str, result, **_kwargs) -> None:
    """Convert the markdown essay into a formatted DOCX document."""
    from docx import Document
    from docx.enum.text import WD_ALIGN_PARAGRAPH

    doc = Document()

    # Title
    title_para = doc.add_heading(query, level=0)
    title_para.alignment = WD_ALIGN_PARAGRAPH.CENTER

    # Parse markdown line by line
    for line in result.essay.split("\n"):
        stripped = line.strip()
        if not stripped:
            doc.add_paragraph("")
            continue

        # Headings
        if stripped.startswith("#### "):
            doc.add_heading(stripped[5:], level=4)
        elif stripped.startswith("### "):
            doc.add_heading(stripped[4:], level=3)
        elif stripped.startswith("## "):
            doc.add_heading(stripped[3:], level=2)
        elif stripped.startswith("# "):
            doc.add_heading(stripped[2:], level=1)
        elif stripped.startswith("- ") or stripped.startswith("* "):
            doc.add_paragraph(stripped[2:], style="List Bullet")
        elif re.match(r"^\d+\.\s", stripped):
            text = re.sub(r"^\d+\.\s", "", stripped)
            doc.add_paragraph(text, style="List Number")
        else:
            doc.add_paragraph(stripped)

    # Sources appendix
    if result.results:
        doc.add_page_break()
        doc.add_heading("Sources", level=1)
        for i, r in enumerate(result.results, 1):
            title = r.title or r.url
            para = doc.add_paragraph(f"{i}. ", style="List Number")
            _add_hyperlink(para, r.url, title)

    doc.save(path)


_WRITERS = {
    "json": _write_json,
    "md": _write_markdown,
    "docx": _write_docx,
}


def _progress_callback(quiet: bool):
    """Return a progress callback (or None if quiet)."""
    if quiet:
        return None
    icons = {
        "start": "🎯",
        "outline": "📋",
        "pass": "🔄",
        "search": "🔍",
        "synthesis": "✍️",
        "complete": "✅",
    }

    def show_progress(stage: str, message: str) -> None:
        icon = icons.get(stage, "•")
        print(f"{icon} {message}", file=sys.stderr)

    return show_progress


async def run_research(
    query: str,
    fmt: str,
    output: str | None,
    preset: str,
    stages: int | None,
    passes: int | None,
    quiet: bool,
) -> None:
    from artemis.config import get_settings
    from artemis.researcher import deep_research

    settings = get_settings()

    # Resolve preset defaults
    if preset == "shallow":
        stages = stages or settings.shallow_research_stages
        passes = passes or settings.shallow_research_passes
        sub_queries = settings.shallow_research_subqueries
        results_per_query = settings.shallow_research_results_per_query
        max_tokens = settings.shallow_research_max_tokens
        content_extraction = settings.shallow_research_content_extraction
        pages_per_section = settings.shallow_research_pages_per_section
        content_max_chars = settings.shallow_research_content_max_chars
    else:
        stages = stages or settings.deep_research_stages
        passes = passes or settings.deep_research_passes
        sub_queries = settings.deep_research_subqueries
        results_per_query = settings.deep_research_results_per_query
        max_tokens = settings.deep_research_max_tokens
        content_extraction = settings.deep_research_content_extraction
        pages_per_section = settings.deep_research_pages_per_section
        content_max_chars = settings.deep_research_content_max_chars

    stdout = output == "-"

    # Resolve output path
    if output is None:
        output = _default_output_path(query, fmt)
    elif stdout and fmt == "docx":
        print("Error: DOCX format cannot be written to stdout.", file=sys.stderr)
        sys.exit(1)

    progress = _progress_callback(quiet or stdout)

    if not quiet and not stdout:
        print(f"Research: {query}", file=sys.stderr)
        print(f"Preset: {preset} | Stages: {stages} | Passes: {passes}", file=sys.stderr)
        print("-" * 50, file=sys.stderr)

    try:
        result = await deep_research(
            query=query,
            stages=stages,
            passes=passes,
            sub_queries_per_stage=sub_queries,
            results_per_query=results_per_query,
            max_tokens=max_tokens,
            content_extraction=content_extraction,
            pages_per_section=pages_per_section,
            content_max_chars=content_max_chars,
            progress_callback=progress,
        )
    except Exception as e:
        print(f"Error: {e}", file=sys.stderr)
        sys.exit(1)

    # Write output
    writer = _WRITERS[fmt]
    if stdout:
        writer(output, query, result, stdout=True)
    else:
        writer(output, query, result)
        if not quiet:
            print(f"\n✅ Saved to {output}", file=sys.stderr)
            print(
                f"   {len(result.essay)} chars | {len(result.results)} sources",
                file=sys.stderr,
            )
            if result.usage:
                print(
                    f"   Tokens: {result.usage.total_tokens} "
                    f"(in={result.usage.input_tokens} out={result.usage.output_tokens})",
                    file=sys.stderr,
                )


def main() -> None:
    parser = argparse.ArgumentParser(
        description="Artemis — one-shot deep research from the command line",
        formatter_class=argparse.RawDescriptionHelpFormatter,
        epilog=(
            "examples:\n"
            '  %(prog)s "quantum computing"\n'
            '  %(prog)s "climate change" --format docx\n'
            '  %(prog)s "rust vs go" --format md -o comparison.md --preset shallow\n'
            '  %(prog)s "LLM architectures" --format md -o -    # stdout\n'
        ),
    )
    parser.add_argument("query", help="Research query or question")
    parser.add_argument(
        "--format", "-f",
        choices=_FORMATS,
        default="json",
        help="Output format (default: json)",
    )
    parser.add_argument(
        "--output", "-o",
        metavar="PATH",
        default=None,
        help='Output file path (default: auto-generated). Use "-" for stdout.',
    )
    parser.add_argument(
        "--preset",
        choices=_PRESETS,
        default="deep",
        help="Research preset (default: deep)",
    )
    parser.add_argument(
        "--stages",
        type=int,
        default=None,
        help="Number of outline sections (overrides preset)",
    )
    parser.add_argument(
        "--passes",
        type=int,
        default=None,
        help="Number of research passes (overrides preset)",
    )
    parser.add_argument(
        "--quiet", "-q",
        action="store_true",
        help="Suppress progress output",
    )
    args = parser.parse_args()

    asyncio.run(
        run_research(
            query=args.query,
            fmt=args.format,
            output=args.output,
            preset=args.preset,
            stages=args.stages,
            passes=args.passes,
            quiet=args.quiet,
        )
    )


if __name__ == "__main__":
    main()
